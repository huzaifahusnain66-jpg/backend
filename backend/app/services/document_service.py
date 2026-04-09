"""
Document Generation Service.
12 premium templates with enhanced cover page, TOC, headers/footers,
drop-shadow images, and styled references for both PDF and DOCX output.
"""

from __future__ import annotations

import os
from datetime import datetime
from typing import List

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.colors import HexColor, white, black, Color
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
    PageBreak, HRFlowable, Table, TableStyle, KeepTogether,
)
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.platypus.flowables import Flowable

from app.services.text_generation_service import GeneratedContent, GeneratedSection
from app.services.image_generation_service import GeneratedImage
from app.utils.file_helpers import ensure_directory, generate_unique_filename
from app.utils.logger import get_logger

logger = get_logger(__name__)

# ── Template Definitions (12 templates) ──────────────────────────────────────
TEMPLATES = {
    # ── 1. Professional (navy blue corporate) ──
    "professional": {
        "primary":     "#003366",
        "secondary":   "#0066CC",
        "accent":      "#E8F0FE",
        "text":        "#1A1A2E",
        "light":       "#F0F4FF",
        "cover_bar":   "#003366",
        "sidebar":     "#E8F0FE",
        "highlight":   "#0066CC",
        "heading_rgb": RGBColor(0, 51, 102),
        "accent_rgb":  RGBColor(0, 102, 204),
        "font":        "Calibri",
        "h_size": 15,  "b_size": 11,
    },
    # ── 2. Academic (dark slate scholarly) ──
    "academic": {
        "primary":     "#2C3E50",
        "secondary":   "#8B0000",
        "accent":      "#FDF6E3",
        "text":        "#2C3E50",
        "light":       "#F8F4EF",
        "cover_bar":   "#2C3E50",
        "sidebar":     "#FDF6E3",
        "highlight":   "#8B0000",
        "heading_rgb": RGBColor(44, 62, 80),
        "accent_rgb":  RGBColor(139, 0, 0),
        "font":        "Times New Roman",
        "h_size": 15,  "b_size": 12,
    },
    # ── 3. Modern (purple/pink bold) ──
    "modern": {
        "primary":     "#6C63FF",
        "secondary":   "#FF6584",
        "accent":      "#F3F0FF",
        "text":        "#2D2D2D",
        "light":       "#F8F7FF",
        "cover_bar":   "#6C63FF",
        "sidebar":     "#F3F0FF",
        "highlight":   "#FF6584",
        "heading_rgb": RGBColor(108, 99, 255),
        "accent_rgb":  RGBColor(255, 101, 132),
        "font":        "Calibri",
        "h_size": 16,  "b_size": 11,
    },
    # ── 4. Minimal (near-black ultra clean) ──
    "minimal": {
        "primary":     "#1A1A1A",
        "secondary":   "#555555",
        "accent":      "#F5F5F5",
        "text":        "#1A1A1A",
        "light":       "#FAFAFA",
        "cover_bar":   "#1A1A1A",
        "sidebar":     "#F5F5F5",
        "highlight":   "#555555",
        "heading_rgb": RGBColor(26, 26, 26),
        "accent_rgb":  RGBColor(85, 85, 85),
        "font":        "Calibri",
        "h_size": 14,  "b_size": 11,
    },
    # ── 5. Colorful (purple/magenta vibrant) ──
    "colorful": {
        "primary":     "#7B2D8B",
        "secondary":   "#E91E8C",
        "accent":      "#FFF0FB",
        "text":        "#2D1B33",
        "light":       "#FDF5FF",
        "cover_bar":   "#7B2D8B",
        "sidebar":     "#FFF0FB",
        "highlight":   "#E91E8C",
        "heading_rgb": RGBColor(123, 45, 139),
        "accent_rgb":  RGBColor(233, 30, 140),
        "font":        "Calibri",
        "h_size": 16,  "b_size": 11,
    },
    # ── 6. Dark (dark mode, cyan accent) ──
    "dark": {
        "primary":     "#0F0F0F",
        "secondary":   "#00D4FF",
        "accent":      "#0A1A1F",
        "text":        "#E0E0E0",
        "light":       "#111111",
        "cover_bar":   "#0F0F0F",
        "sidebar":     "#0A1A1F",
        "highlight":   "#00D4FF",
        "heading_rgb": RGBColor(0, 212, 255),
        "accent_rgb":  RGBColor(0, 212, 255),
        "font":        "Calibri",
        "h_size": 15,  "b_size": 11,
    },
    # ── 7. Nature (deep green forest) ──
    "nature": {
        "primary":     "#1B5E20",
        "secondary":   "#4CAF50",
        "accent":      "#F1F8E9",
        "text":        "#1B3A1E",
        "light":       "#F9FBF2",
        "cover_bar":   "#1B5E20",
        "sidebar":     "#F1F8E9",
        "highlight":   "#8BC34A",
        "heading_rgb": RGBColor(27, 94, 32),
        "accent_rgb":  RGBColor(76, 175, 80),
        "font":        "Calibri",
        "h_size": 15,  "b_size": 11,
    },
    # ── 8. Sunset (orange/amber warm) ──
    "sunset": {
        "primary":     "#BF360C",
        "secondary":   "#FF6F00",
        "accent":      "#FFF8E1",
        "text":        "#3E2000",
        "light":       "#FFFDE7",
        "cover_bar":   "#BF360C",
        "sidebar":     "#FFF8E1",
        "highlight":   "#FFD54F",
        "heading_rgb": RGBColor(191, 54, 12),
        "accent_rgb":  RGBColor(255, 111, 0),
        "font":        "Calibri",
        "h_size": 15,  "b_size": 11,
    },
    # ── 9. Ocean (deep sea blue) ──
    "ocean": {
        "primary":     "#01579B",
        "secondary":   "#00B0FF",
        "accent":      "#E0F7FA",
        "text":        "#01294E",
        "light":       "#F0FBFF",
        "cover_bar":   "#01579B",
        "sidebar":     "#E0F7FA",
        "highlight":   "#E0F7FA",
        "heading_rgb": RGBColor(1, 87, 155),
        "accent_rgb":  RGBColor(0, 176, 255),
        "font":        "Calibri",
        "h_size": 15,  "b_size": 11,
    },
    # ── 10. Rose (elegant rose pink) ──
    "rose": {
        "primary":     "#880E4F",
        "secondary":   "#E91E63",
        "accent":      "#FCE4EC",
        "text":        "#3E0027",
        "light":       "#FFF0F5",
        "cover_bar":   "#880E4F",
        "sidebar":     "#FCE4EC",
        "highlight":   "#FCE4EC",
        "heading_rgb": RGBColor(136, 14, 79),
        "accent_rgb":  RGBColor(233, 30, 99),
        "font":        "Calibri",
        "h_size": 15,  "b_size": 11,
    },
    # ── 11. Midnight (deep indigo/violet) ──
    "midnight": {
        "primary":     "#1A237E",
        "secondary":   "#7C4DFF",
        "accent":      "#EDE7F6",
        "text":        "#0D1240",
        "light":       "#F3F0FF",
        "cover_bar":   "#1A237E",
        "sidebar":     "#EDE7F6",
        "highlight":   "#B388FF",
        "heading_rgb": RGBColor(26, 35, 126),
        "accent_rgb":  RGBColor(124, 77, 255),
        "font":        "Calibri",
        "h_size": 15,  "b_size": 11,
    },
    # ── 12. Corporate (black/red executive) ──
    "corporate": {
        "primary":     "#212121",
        "secondary":   "#B71C1C",
        "accent":      "#FFEBEE",
        "text":        "#212121",
        "light":       "#FAFAFA",
        "cover_bar":   "#212121",
        "sidebar":     "#FFEBEE",
        "highlight":   "#FF5252",
        "heading_rgb": RGBColor(33, 33, 33),
        "accent_rgb":  RGBColor(183, 28, 28),
        "font":        "Calibri",
        "h_size": 15,  "b_size": 11,
    },
}


# ── Custom ReportLab Flowables ────────────────────────────────────────────────

class GradientRect(Flowable):
    """Horizontal gradient bar rendered as 80 thin vertical stripes."""

    def __init__(self, width, height, color1: HexColor, color2: HexColor):
        super().__init__()
        self.width  = width
        self.height = height
        self.color1 = color1
        self.color2 = color2

    def _hex_to_rgb(self, color: HexColor):
        return color.red, color.green, color.blue

    def draw(self):
        stripes = 80
        stripe_w = self.width / stripes
        r1, g1, b1 = self._hex_to_rgb(self.color1)
        r2, g2, b2 = self._hex_to_rgb(self.color2)
        for i in range(stripes):
            t = i / max(stripes - 1, 1)
            r = r1 + (r2 - r1) * t
            g = g1 + (g2 - g1) * t
            b = b1 + (b2 - b1) * t
            self.canv.setFillColorRGB(r, g, b)
            self.canv.rect(i * stripe_w, 0, stripe_w + 0.5, self.height, fill=1, stroke=0)


class ColorRect(Flowable):
    """Solid filled rectangle."""

    def __init__(self, width, height, color):
        super().__init__()
        self.width  = width
        self.height = height
        self.color  = color

    def draw(self):
        self.canv.setFillColor(self.color)
        self.canv.rect(0, 0, self.width, self.height, fill=1, stroke=0)


class ImageFrame(Flowable):
    """
    Image with drop shadow, white padding frame, colored border, and caption bar.
    Gracefully skips if the image path is missing or unreadable.
    """

    def __init__(self, img_path: str, caption: str, width: float, height: float,
                 accent_color: HexColor, fn_italic: str = "Helvetica-Oblique"):
        super().__init__()
        self.img_path    = img_path
        self.caption     = caption
        self.img_width   = width
        self.img_height  = height
        self.accent      = accent_color
        self.fn_italic   = fn_italic
        self.padding     = 6
        self.shadow_off  = 3
        self.cap_h       = 18
        # Total flowable dimensions (shadow adds to bottom-right)
        self.width  = width  + self.padding * 2 + self.shadow_off
        self.height = height + self.padding * 2 + self.cap_h + self.shadow_off

    def draw(self):
        c = self.canv
        p = self.padding
        so = self.shadow_off
        cap_h = self.cap_h

        # Drop shadow
        c.setFillColorRGB(0.5, 0.5, 0.5, 0.3)
        c.rect(so, -so, self.img_width + p * 2, self.img_height + p * 2 + cap_h,
               fill=1, stroke=0)

        # White padding frame
        c.setFillColor(white)
        c.rect(0, 0, self.img_width + p * 2, self.img_height + p * 2 + cap_h,
               fill=1, stroke=0)

        # Colored border (1.5pt)
        c.setStrokeColor(self.accent)
        c.setLineWidth(1.5)
        c.rect(0, 0, self.img_width + p * 2, self.img_height + p * 2 + cap_h,
               fill=0, stroke=1)

        # Caption bar at bottom
        c.setFillColor(self.accent)
        c.rect(0, 0, self.img_width + p * 2, cap_h, fill=1, stroke=0)

        # Caption text
        c.setFillColor(white)
        c.setFont(self.fn_italic, 8)
        c.drawCentredString((self.img_width + p * 2) / 2, 5, self.caption[:90])

        # Image itself
        try:
            c.drawImage(
                self.img_path,
                p, cap_h + p,
                width=self.img_width, height=self.img_height,
                preserveAspectRatio=True, mask="auto",
            )
        except Exception as e:
            logger.warning("ImageFrame: could not draw image '%s' | %s", self.img_path, e)
            c.setFillColor(HexColor("#EEEEEE"))
            c.rect(p, cap_h + p, self.img_width, self.img_height, fill=1, stroke=0)
            c.setFillColor(HexColor("#999999"))
            c.setFont("Helvetica", 9)
            c.drawCentredString(p + self.img_width / 2, cap_h + p + self.img_height / 2,
                                "[Image unavailable]")


# ── Document Service ──────────────────────────────────────────────────────────

class DocumentService:
    def __init__(self, storage_path: str = "storage/documents") -> None:
        self._storage_path = os.path.abspath(storage_path)
        ensure_directory(self._storage_path)

    # ══════════════════════════════════════════════════════════════════════════
    #  DOCX
    # ══════════════════════════════════════════════════════════════════════════

    def generate_docx(
        self,
        content: GeneratedContent,
        images: List[GeneratedImage],
        template: str = "professional",
        layout: str = "standard",
        assignment_id: str = "",
        student_name: str = "",
        roll_number: str = "",
        department: str = "",
    ) -> str:
        tmpl = TEMPLATES.get(template, TEMPLATES["professional"])
        doc  = Document()

        # ── Page margins ──
        for sec in doc.sections:
            sec.top_margin    = Cm(2.5)
            sec.bottom_margin = Cm(2.5)
            sec.left_margin   = Cm(3)
            sec.right_margin  = Cm(2.5)

        # ── Default style ──
        normal = doc.styles["Normal"]
        _font  = tmpl["font"] if tmpl["font"] != "Times New Roman" else "Times New Roman"
        normal.font.name                       = _font
        normal.font.size                       = Pt(tmpl["b_size"])
        normal.paragraph_format.line_spacing   = 1.5
        normal.paragraph_format.space_after    = Pt(6)
        normal.paragraph_format.alignment      = WD_ALIGN_PARAGRAPH.JUSTIFY

        # ── Cover page ──
        self._docx_cover(doc, content.title, student_name, roll_number, department, tmpl)
        doc.add_page_break()

        # ── Introduction ──
        self._docx_heading(doc, "Introduction", tmpl)
        self._docx_body(doc, content.introduction, tmpl)

        # ── Body sections ──
        image_map = {img.section_title: img for img in images if img.success}
        for section in content.sections:
            self._docx_heading(doc, section.title, tmpl)
            self._docx_body(doc, section.content, tmpl)
            img = image_map.get(section.title)
            if img and img.image_path and os.path.exists(img.image_path):
                self._docx_image(doc, img, tmpl)

        # ── Conclusion ──
        self._docx_heading(doc, "Conclusion", tmpl)
        self._docx_body(doc, content.conclusion, tmpl)

        # ── References ──
        doc.add_page_break()
        self._docx_heading(doc, "References", tmpl)
        for i, ref in enumerate(content.references, 1):
            rp = doc.add_paragraph()
            rp.paragraph_format.left_indent       = Inches(0.5)
            rp.paragraph_format.first_line_indent = Inches(-0.5)
            rp.paragraph_format.space_after       = Pt(4)
            run = rp.add_run(f"{i}. {ref}")
            run.font.size = Pt(10)
            run.font.name = tmpl["font"]

        filename = generate_unique_filename("docx", prefix=f"assignment_{assignment_id[:8]}")
        filepath = os.path.join(self._storage_path, filename)
        doc.save(filepath)
        logger.info("DOCX generated | %s", filepath)
        return filepath

    # ── DOCX helpers ──────────────────────────────────────────────────────────

    def _docx_cover(self, doc, title, student_name, roll_number, department, tmpl):
        """Premium DOCX cover page with shaded bars and centered info."""
        font_name = tmpl["font"]

        # Top colored bar (tall, 32pt font space)
        bar_top = doc.add_paragraph()
        bar_top.paragraph_format.space_before = Pt(0)
        bar_top.paragraph_format.space_after  = Pt(0)
        self._set_para_shading(bar_top, tmpl["cover_bar"])
        bar_top.add_run("  ").font.size = Pt(32)

        doc.add_paragraph("")

        # Eyebrow text
        eyebrow = doc.add_paragraph()
        eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = eyebrow.add_run("ASSIGNMENT SUBMISSION")
        r.font.size      = Pt(11)
        r.font.color.rgb = tmpl["accent_rgb"]
        r.font.bold      = True
        r.font.name      = font_name

        doc.add_paragraph("")

        # Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(10)
        tr = title_p.add_run(title)
        tr.bold           = True
        tr.font.size      = Pt(24)
        tr.font.color.rgb = tmpl["heading_rgb"]
        tr.font.name      = font_name

        # Decorative divider
        div = doc.add_paragraph()
        div.alignment = WD_ALIGN_PARAGRAPH.CENTER
        div.paragraph_format.space_before = Pt(8)
        div.paragraph_format.space_after  = Pt(8)
        dr = div.add_run("━" * 40)
        dr.font.color.rgb = tmpl["accent_rgb"]
        dr.font.size      = Pt(12)
        dr.font.name      = font_name

        doc.add_paragraph("")

        # Student info: label bold in heading color | value normal
        rows: list[tuple[str, str]] = []
        if student_name: rows.append(("Student Name",    student_name))
        if roll_number:  rows.append(("Roll Number",     roll_number))
        if department:   rows.append(("Department",      department))
        rows.append(("Submission Date", datetime.now().strftime("%B %d, %Y")))

        for label, value in rows:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            lb = p.add_run(f"{label}:  ")
            lb.bold           = True
            lb.font.size      = Pt(12)
            lb.font.color.rgb = tmpl["heading_rgb"]
            lb.font.name      = font_name
            vr = p.add_run(value)
            vr.font.size = Pt(12)
            vr.font.name = font_name

        doc.add_paragraph("")

        # Bottom colored bar
        bot = doc.add_paragraph()
        bot.paragraph_format.space_before = Pt(20)
        bot.paragraph_format.space_after  = Pt(0)
        self._set_para_shading(bot, tmpl["cover_bar"])
        bot.add_run("  ").font.size = Pt(14)

    def _set_para_shading(self, para, hex_color: str):
        """Set paragraph background color via XML shading element."""
        hex_color = hex_color.lstrip("#")
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        pPr.append(shd)

    def _docx_heading(self, doc, text, tmpl):
        """Section heading: Heading 1 style, template font/color, 14pt before / 6pt after."""
        h = doc.add_heading(level=1)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after  = Pt(6)
        run = h.add_run(text)
        run.font.size      = Pt(tmpl["h_size"])
        run.font.color.rgb = tmpl["heading_rgb"]
        run.font.bold      = True
        run.font.name      = tmpl["font"]

    def _docx_body(self, doc, text: str, tmpl):
        """Justified body paragraphs with correct font and 1.5 line spacing."""
        for para in text.split("\n\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(para)
                run.font.name = tmpl["font"]
                run.font.size = Pt(tmpl["b_size"])

    def _docx_image(self, doc, img: GeneratedImage, tmpl):
        """Insert centered image (5.5 in) with italic caption in accent color below."""
        try:
            doc.add_paragraph("")
            pic_para = doc.add_paragraph()
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_para.add_run()
            run.add_picture(img.image_path, width=Inches(5.5))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(img.caption)
            cr.italic         = True
            cr.font.size      = Pt(9)
            cr.font.color.rgb = tmpl["accent_rgb"]
            cr.font.name      = tmpl["font"]
            doc.add_paragraph("")
        except Exception as e:
            logger.warning("DOCX image insert failed | %s", e)

    # ══════════════════════════════════════════════════════════════════════════
    #  PDF
    # ══════════════════════════════════════════════════════════════════════════

    def generate_pdf(
        self,
        content: GeneratedContent,
        images: List[GeneratedImage],
        template: str = "professional",
        layout: str = "standard",
        assignment_id: str = "",
        student_name: str = "",
        roll_number: str = "",
        department: str = "",
    ) -> str:
        tmpl = TEMPLATES.get(template, TEMPLATES["professional"])
        filename = generate_unique_filename("pdf", prefix=f"assignment_{assignment_id[:8]}")
        filepath = os.path.join(self._storage_path, filename)

        primary      = HexColor(tmpl["primary"])
        secondary    = HexColor(tmpl["secondary"])
        text_color   = HexColor(tmpl["text"])
        light_color  = HexColor(tmpl["light"])
        accent_color = HexColor(tmpl["accent"])
        highlight    = HexColor(tmpl["highlight"])

        # ── Font selection ──
        if tmpl["font"] == "Times New Roman":
            fn        = "Times-Roman"
            fn_bold   = "Times-Bold"
            fn_italic = "Times-Italic"
        else:
            fn        = "Helvetica"
            fn_bold   = "Helvetica-Bold"
            fn_italic = "Helvetica-Oblique"

        W, H   = A4
        margin = 72  # 1 inch
        body_w = W - margin * 2

        # ── Paragraph Styles ──
        cover_eyebrow = ParagraphStyle(
            "CoverEyebrow", fontName=fn_bold, fontSize=9,
            textColor=secondary, alignment=TA_CENTER, spaceAfter=6, leading=14,
        )
        cover_title = ParagraphStyle(
            "CoverTitle", fontName=fn_bold, fontSize=28,
            textColor=primary, alignment=TA_CENTER, spaceAfter=10, leading=36,
        )
        toc_title_style = ParagraphStyle(
            "TOCTitle", fontName=fn_bold, fontSize=11,
            textColor=text_color, alignment=TA_LEFT, leading=16,
        )
        h_style = ParagraphStyle(
            "Heading", fontName=fn_bold, fontSize=tmpl["h_size"],
            textColor=primary, spaceBefore=16, spaceAfter=6, leading=22,
        )
        body_style = ParagraphStyle(
            "Body", fontName=fn, fontSize=tmpl["b_size"],
            textColor=text_color, leading=18, alignment=TA_JUSTIFY, spaceAfter=8,
        )
        caption_style = ParagraphStyle(
            "Caption", fontName=fn_italic, fontSize=9,
            textColor=secondary, alignment=TA_CENTER, spaceAfter=12,
        )
        ref_num_style = ParagraphStyle(
            "RefNum", fontName=fn_bold, fontSize=10,
            textColor=secondary, alignment=TA_LEFT, leading=14,
        )
        ref_style = ParagraphStyle(
            "Ref", fontName=fn, fontSize=10,
            textColor=text_color, leading=14,
            leftIndent=0, spaceAfter=5,
        )

        story = []

        # ══════════════════════════════════════════════════════════════════
        #  COVER PAGE
        # ══════════════════════════════════════════════════════════════════

        # Top gradient bar — full body width, 18pt high
        story.append(GradientRect(body_w, 18, primary, secondary))
        story.append(Spacer(1, 0.45 * inch))

        # Eyebrow
        eyebrow_p = Paragraph("ACADEMIC ASSIGNMENT", cover_eyebrow)
        story.append(eyebrow_p)
        story.append(Spacer(1, 0.15 * inch))

        # Title
        story.append(Paragraph(content.title, cover_title))
        story.append(Spacer(1, 0.05 * inch))

        # Horizontal rule — 60% width
        story.append(HRFlowable(
            width="60%", thickness=2, color=secondary,
            spaceAfter=16, spaceBefore=4,
        ))
        story.append(Spacer(1, 0.25 * inch))

        # Student info as 2-column table: label RIGHT bold | value LEFT normal
        info_rows: list[tuple[str, str]] = []
        if student_name: info_rows.append(("Student Name",    student_name))
        if roll_number:  info_rows.append(("Roll Number",     roll_number))
        if department:   info_rows.append(("Department",      department))
        info_rows.append(("Submission Date", datetime.now().strftime("%B %d, %Y")))

        for label, value in info_rows:
            tbl_data = [[
                Paragraph(
                    f"<b>{label}</b>",
                    ParagraphStyle("tl", fontName=fn_bold, fontSize=9,
                                   textColor=secondary, alignment=TA_RIGHT),
                ),
                Paragraph(
                    value,
                    ParagraphStyle("tv", fontName=fn, fontSize=11,
                                   textColor=text_color, alignment=TA_LEFT),
                ),
            ]]
            tbl = Table(tbl_data, colWidths=[2.5 * inch, 3.5 * inch])
            tbl.setStyle(TableStyle([
                ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING",    (0, 0), (-1, -1), 6),
                ("LINEBELOW",     (0, 0), (-1, -1), 0.5, HexColor("#DDDDDD")),
            ]))
            story.append(tbl)

        story.append(Spacer(1, 0.5 * inch))

        # Bottom gradient bar (secondary → primary)
        story.append(GradientRect(body_w, 12, secondary, primary))
        story.append(PageBreak())

        # ══════════════════════════════════════════════════════════════════
        #  TABLE OF CONTENTS (page 2)
        # ══════════════════════════════════════════════════════════════════

        # Full-width colored header box
        story.append(ColorRect(body_w, 28, primary))
        story.append(Spacer(1, -28))     # overlap so text sits on the bar
        toc_header_p = Paragraph(
            "<font color='white'><b>TABLE OF CONTENTS</b></font>",
            ParagraphStyle("TOCHeader", fontName=fn_bold, fontSize=13,
                           textColor=white, alignment=TA_CENTER, leading=28),
        )
        story.append(toc_header_p)
        story.append(Spacer(1, 10))

        # Build TOC entries
        toc_sections = ["Introduction"] + [s.title for s in content.sections] + ["Conclusion", "References"]
        for idx, sec_title in enumerate(toc_sections, 1):
            row_bg = HexColor("#F9F9F9") if idx % 2 == 0 else white
            num_cell = Paragraph(
                f"<b>{idx}</b>",
                ParagraphStyle("TOCNum", fontName=fn_bold, fontSize=11,
                               textColor=secondary, alignment=TA_CENTER, leading=16),
            )
            title_cell = Paragraph(
                sec_title,
                ParagraphStyle("TOCEntry", fontName=fn, fontSize=10,
                               textColor=text_color, alignment=TA_LEFT, leading=16),
            )
            dots_cell = Paragraph(
                "." * 50,
                ParagraphStyle("TOCDots", fontName=fn, fontSize=9,
                               textColor=HexColor("#AAAAAA"), alignment=TA_LEFT, leading=16),
            )
            page_cell = Paragraph(
                f"<b>{idx + 1}</b>",
                ParagraphStyle("TOCPage", fontName=fn_bold, fontSize=10,
                               textColor=primary, alignment=TA_RIGHT, leading=16),
            )
            toc_row_data = [[num_cell, title_cell, dots_cell, page_cell]]
            toc_tbl = Table(toc_row_data, colWidths=[0.4 * inch, 3.0 * inch, 2.0 * inch, 0.5 * inch])
            toc_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (-1, -1), row_bg),
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING",    (0, 0), (-1, -1), 7),
                ("LINEBELOW",     (0, 0), (-1, -1), 0.5, HexColor("#E0E0E0")),
            ]))
            story.append(toc_tbl)

        story.append(PageBreak())

        # ══════════════════════════════════════════════════════════════════
        #  CONTENT SECTIONS
        # ══════════════════════════════════════════════════════════════════

        image_map = {img.section_title: img for img in images if img.success}
        section_number = [0]  # mutable counter

        def add_section(title: str, text: str, img=None):
            section_number[0] += 1
            num_str = str(section_number[0])

            # Section header: two-column layout — number LEFT | title RIGHT
            num_cell = Paragraph(
                f"<font color='{tmpl['secondary']}'><b>{num_str}</b></font>",
                ParagraphStyle("SecNum", fontName=fn_bold, fontSize=20,
                               textColor=secondary, alignment=TA_CENTER, leading=28),
            )
            title_cell = Paragraph(
                title.upper(),
                ParagraphStyle("SecTitle", fontName=fn_bold, fontSize=tmpl["h_size"],
                               textColor=primary, alignment=TA_LEFT, leading=22),
            )
            hdr_data = [[num_cell, title_cell]]
            hdr_tbl = Table(hdr_data, colWidths=[0.6 * inch, body_w - 0.6 * inch])
            hdr_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (0, 0), light_color),
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING",    (0, 0), (-1, -1), 8),
                ("LINEBELOW",     (0, 0), (-1, -1), 2, secondary),
            ]))
            story.append(hdr_tbl)
            story.append(Spacer(1, 6))

            # Body paragraphs
            for para in text.split("\n\n"):
                para = para.strip()
                if para:
                    story.append(Paragraph(para, body_style))

            # Image with ImageFrame flowable
            if img and img.image_path and os.path.exists(img.image_path):
                try:
                    frm = ImageFrame(
                        img_path     = img.image_path,
                        caption      = img.caption,
                        width        = 4.8 * inch,
                        height       = 3.0 * inch,
                        accent_color = secondary,
                        fn_italic    = fn_italic,
                    )
                    frm.hAlign = "CENTER"
                    story.append(KeepTogether([Spacer(1, 8), frm]))
                except Exception as e:
                    logger.warning("PDF ImageFrame failed | %s", e)

            story.append(Spacer(1, 10))

        add_section("Introduction", content.introduction)

        for section in content.sections:
            add_section(section.title, section.content, image_map.get(section.title))

        add_section("Conclusion", content.conclusion)

        # ══════════════════════════════════════════════════════════════════
        #  REFERENCES PAGE
        # ══════════════════════════════════════════════════════════════════

        story.append(PageBreak())

        # Full-width primary colored header box
        story.append(ColorRect(body_w, 28, primary))
        story.append(Spacer(1, -28))
        story.append(Paragraph(
            "<font color='white'><b>REFERENCES</b></font>",
            ParagraphStyle("RefHeader", fontName=fn_bold, fontSize=13,
                           textColor=white, alignment=TA_CENTER, leading=28),
        ))
        story.append(Spacer(1, 10))

        for i, ref in enumerate(content.references, 1):
            ref_data = [[
                Paragraph(
                    f"<b>[{i}]</b>",
                    ParagraphStyle("rnum", fontName=fn_bold, fontSize=10,
                                   textColor=secondary, alignment=TA_LEFT, leading=14),
                ),
                Paragraph(
                    ref,
                    ParagraphStyle("rtxt", fontName=fn, fontSize=10,
                                   textColor=text_color, leading=14, alignment=TA_LEFT),
                ),
            ]]
            ref_tbl = Table(ref_data, colWidths=[0.55 * inch, body_w - 0.55 * inch])
            ref_tbl.setStyle(TableStyle([
                ("VALIGN",        (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING",    (0, 0), (-1, -1), 5),
                ("LINEBELOW",     (0, 0), (-1, -1), 0.5, HexColor("#DDDDDD")),
            ]))
            story.append(ref_tbl)

        # ══════════════════════════════════════════════════════════════════
        #  BUILD WITH HEADER / FOOTER CALLBACKS
        # ══════════════════════════════════════════════════════════════════

        doc_pdf = SimpleDocTemplate(
            filepath, pagesize=A4,
            rightMargin=margin, leftMargin=margin,
            topMargin=margin,   bottomMargin=margin,
        )

        title_short = content.title[:55]

        def _on_first_page(canvas, doc):
            """Cover page — no header, no footer."""
            pass

        def _on_later_pages(canvas, doc):
            """Content pages: top header bar + bottom footer bar with page circle."""
            canvas.saveState()

            # ── TOP HEADER BAR ──
            # 3px primary color line at very top
            canvas.setStrokeColor(primary)
            canvas.setLineWidth(3)
            canvas.line(margin, H - 36, W - margin, H - 36)

            # Light background strip
            canvas.setFillColor(HexColor("#F7F7F7"))
            canvas.rect(margin, H - 54, W - margin * 2, 17, fill=1, stroke=0)

            # Title LEFT
            canvas.setFont(fn_bold, 7.5)
            canvas.setFillColor(primary)
            canvas.drawString(margin + 4, H - 49, title_short)

            # Date RIGHT
            canvas.setFont(fn, 7.5)
            canvas.setFillColor(HexColor("#666666"))
            date_str = datetime.now().strftime("%B %d, %Y")
            canvas.drawRightString(W - margin - 4, H - 49, date_str)

            # ── BOTTOM FOOTER BAR ──
            # 2px primary color top border
            canvas.setStrokeColor(primary)
            canvas.setLineWidth(2)
            canvas.line(margin, 54, W - margin, 54)

            # Light background strip
            canvas.setFillColor(HexColor("#F7F7F7"))
            canvas.rect(margin, 36, W - margin * 2, 17, fill=1, stroke=0)

            # "AI Assignment Generator" LEFT
            canvas.setFont(fn, 7.5)
            canvas.setFillColor(HexColor("#888888"))
            canvas.drawString(margin + 4, 41, "AI Assignment Generator")

            # Page number circle RIGHT
            page_num = doc.page - 1  # cover = page 1, so content starts at 1
            cx = W - margin - 12
            cy = 45
            canvas.setFillColor(primary)
            canvas.circle(cx, cy, 10, fill=1, stroke=0)
            canvas.setFillColor(white)
            canvas.setFont(fn_bold, 8)
            canvas.drawCentredString(cx, cy - 3, str(page_num))

            canvas.restoreState()

        doc_pdf.build(story, onFirstPage=_on_first_page, onLaterPages=_on_later_pages)
        logger.info("PDF generated | %s", filepath)
        return filepath
